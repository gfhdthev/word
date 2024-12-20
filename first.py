import docx

doc = docx.Document('first.docx')

sentence = int(input('Введите, в какое предложение вам надо вставить слово: '))

k = 0 #переменная для подсчета
start = 0 #начало предложения (индекс)
finish = 0 #конец предложения(индекс)
sent = [] #для работы с предложением
symbols = ['.', '!', '?'] #символы конца предложения
flag = False #флаг

for paragraph in doc.paragraphs:
    sent.append(paragraph.text) #добавляем весь текст в лист

sent = '\n'.join(sent) #делаем его одной строкой
#ищем индекс начала предложения, которое ввел пользователь
for word in sent:
    for letter in word: #считаем кол-во букв в промежутке до начала предложения
        if k+1 == sentence: #когда кол-во предложений будет на 1 меньше чем вводимое число, то останавливаемся 
            #(начало и конец предложения считаетмся по знакам из списка symbols)
            flag = True 
            break
        start += 1 #индекс начала предложения
        if letter in symbols:
            k += 1 #кол-во пройденных предложений
    if flag:
        break
#обнуляем переменные
k = 0
flag = False
#тоже самое, только с концом предложения, которое ввел пользователь
for word in sent:
    for letter in word:
        if k == sentence: #только уже проходим до k раное нашему предложению, чтобы точка была его концом
            flag = True
            break
        finish += 1
        if letter in symbols:
            k += 1
    if flag:
        break
#сохраняем наше предложение через срезы 
if sent[start] == ' ':
    find = sent[start+1:finish]
else:
    find = sent[start:finish]

k = 0
#т.к. предложения ищу через значи препинания, то мы могли случайно захватить какое-то оглавление, проверяем символы на перенос строки('\n')
#и если он есть, то обрезаем строку после выреза
if '\n' in find:
    print(find.index('\n'))
    find = find[find.index('\n')+1:]

'''Вставляем этот код 2 раза, т.к. программа считает '\n' перед заголовком (в части сверху удаляли как раз его)'''
#в данной части удаляем непостредственно сам зоголовок
if '\n' in find:
    print(find.index('\n'))
    find = find[find.index('\n')+1:]
    k += 1 #прибаляем к k 1, т.к. в следующей части кода, где мы ищем параграф, оно будет считать найденный параграф за оглавление

'''!!!ERROR!!!'''
#ищем в каком параграфе находится наше предложение
for paragraph in doc.paragraphs:
    if find in paragraph.text:
        found_paragraph = k
        break
    k += 1
'''!!!ERROR!!!'''

print(find)

after_word = input('Введите после какого слова вы хотите вставить: ')
add_word = input('Введите, что вы хотите вставить: ')
find_words = find.split() #разделяем на слова
#надо отделить знаки, если они стоят с какими-то словами
znaki = [',', ':', ';', '"', '(', ')', '.', '!', '?']


for wordd in find_words: #проходим по всем словам
    if wordd[-1] in znaki: #если последний элемент включается в список
        if len(wordd) == 1 and wordd in znaki: #если его длина равна 1, то пропускаем его
            continue
        else:
            ind = find_words.index(wordd) #записываем индекс, куда надо будет вставить слово без знака
            znak = wordd[-1] #записываем сам знак
            find_words[ind] = wordd[:-1] #на место слова со знаком ставим слово без знака
            find_words.insert(ind+1, znak) #и на следующее место возвращаем этот знак

'''!!!ERROR!!!'''
for worrd in find_words:
    if worrd == after_word: #если какое-либо слово из списка совпадает
        find_words.insert((find_words.index(worrd))+1, add_word) #то на следующее место вставляем слово(текст), которое ввел пользователь
        break
'''!!!ERROR!!!'''

all_sentence = ' '.join(find_words) #переделываем в одну строку

'''!!!ERROR!!!'''
#сейчас у нас перед знаками стоят пробелы, их надо удалить
for elem in all_sentence:
    print(elem)
    if elem in znaki:
        ind = all_sentence.index(elem)
        if all_sentence[ind-1] == ' ': #если перед знаком стоит пробел
            print(all_sentence.index(elem))
            first = all_sentence[:ind-1] #то срезаем его
            second = all_sentence[ind:] 
            all_sentence = first + second #и объединяем 
'''!!!ERROR!!!'''

'''!!!ERROR!!!'''
'''
#первое предложение ставится с пробелом
new_par = doc.add_paragraph()
new_par.add_run(doc.paragraphs[found_paragraph].text[:start]) #через срезы создаем в конце списка новый параграф
new_par.add_run(' ')
new_par.add_run(all_sentence) #на место, где мы срезали, вставляем предлоджение, в которое мы добавляли слово
new_par.add_run(doc.paragraphs[found_paragraph].text[finish:])
'''
'''!!!ERROR!!!'''

print(all_sentence)

doc.save('test.docx')